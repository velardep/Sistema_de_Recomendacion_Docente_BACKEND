# app/application/use_cases/ingesta_espacio_archivo.py

# Este use case pertenece principalmente a los flujos de ESPACIOS DE TRABAJO,
# EMBEDDINGS, RED1, RED2 y RED3. Se encarga de procesar un archivo subido a un
# espacio: validar acceso, extraer y limpiar su texto, dividirlo en chunks,
# generar embeddings, guardar esos embeddings en lote, clasificar cada chunk con
# RED1, generar etiquetas de RED2 cuando corresponda y registrar el evento en RED3.

from __future__ import annotations
from fastapi import UploadFile
from typing import List, Dict, Any, Optional

import re
import hashlib
import uuid
import time

# Límites de ingesta: solo se rechaza duro por tamaño bruto, formato inválido
# o falta de texto extraíble. Lo demás se recorta para rescatar el máximo útil.
ALLOWED_EXTENSIONS = (".pdf", ".docx")
MAX_FILE_BYTES = 8 * 1024 * 1024          # 8 MB máximo por archivo (rechazo duro)
MIN_TEXT_CHARS = 500                      # mínimo de texto útil extraído
MAX_TEXT_CHARS = 600_000                  # tope blando de texto procesable
MAX_PDF_PAGES = 120                       # tope blando de páginas PDF a leer
MAX_CHUNKS = 180                          # tope blando de chunks a procesar


class IngestarArchivoEspacioUseCase:
    def __init__(
        self,
        auth_client,
        espacios_repo,
        embeddings_model,
        embeddings_repo,
        red1_service,
        red2_model=None,
        red2_repo=None,
        red3_service=None,  
        espacio_archivos_repo=None,
    ):
        self.auth = auth_client
        self.espacios_repo = espacios_repo
        self.embeddings_model = embeddings_model
        self.embeddings_repo = embeddings_repo
        self.red1 = red1_service
        self.red2_model = red2_model
        self.red2_repo = red2_repo
        self.red3_service = red3_service  
        self.espacio_archivos_repo = espacio_archivos_repo

    async def execute(self, access_token: str, espacio_id: str, file: UploadFile, file_id: str | None = None) -> dict:
        # Primero se identifica al docente autenticado y se valida que el espacio exista
        # y sea accesible antes de iniciar cualquier procesamiento del archivo.
        user = await self.auth.get_user(access_token)
        docente_id = user["id"]

        espacio = await self.espacios_repo.obtener(access_token, espacio_id)
        if not espacio:
            return {"ok": False, "detail": "Espacio no encontrado"}

        # Se obtienen los metadatos básicos del archivo y se carga su contenido completo
        # en memoria para validar que realmente tenga datos útiles.
        filename = (file.filename or "archivo").strip()
        content_type = (file.content_type or "").lower()

        data = await file.read()
        if not data or len(data) < 10:
            return {"ok": False, "detail": "Archivo vacío o inválido"}
        
        # Identificador lógico del archivo y huella para relacionar su procesamiento.
        file_id = file_id or str(uuid.uuid4())
        file_hash = hashlib.sha1(data).hexdigest()

        # Crea el registro maestro del archivo antes del procesamiento para poder listarlo luego.
        if self.espacio_archivos_repo:
            try:
                await self.espacio_archivos_repo.crear(
                    access_token,
                    {
                        "id": file_id,
                        "docente_id": docente_id,
                        "espacio_id": espacio_id,
                        "filename_original": filename,
                        "mime_type": content_type or None,
                        "size_bytes": len(data),
                        "file_hash": file_hash,
                        "total_chunks": 0,
                        "estado": "processing",
                        "error_detail": None,
                    },
                )
            except Exception:
                # No se rompe la ingesta principal por una falla aislada en el registro maestro.
                pass
        
        # Se limita el tamaño bruto para evitar cargas excesivas en memoria y procesamiento.
        if len(data) > MAX_FILE_BYTES:
            if self.espacio_archivos_repo:
                try:
                    await self.espacio_archivos_repo.actualizar(
                        access_token,
                        file_id,
                        {
                            "estado": "failed",
                            "error_detail": "El archivo supera el tamaño máximo permitido de 8 MB.",
                        },
                    )
                except Exception:
                    pass

            return {
                "ok": False,
                "detail": "El archivo supera el tamaño máximo permitido de 8 MB."
            }

        # Aquí se detecta el tipo de archivo y se aplica la estrategia de extracción de
        # texto correspondiente. Si no es PDF ni DOCX, se intenta tratar como texto plano.
        # Se restringe la ingesta únicamente a PDF y DOCX para evitar formatos no controlados.
        text = ""
        lowered_name = filename.lower()

        is_pdf = ("pdf" in content_type) or lowered_name.endswith(".pdf")
        is_docx = (
            ("word" in content_type)
            or ("officedocument.wordprocessingml.document" in content_type)
            or lowered_name.endswith(".docx")
        )

        if not (is_pdf or is_docx):
            if self.espacio_archivos_repo:
                try:
                    await self.espacio_archivos_repo.actualizar(
                        access_token,
                        file_id,
                        {
                            "estado": "failed",
                            "error_detail": "Formato no permitido. Solo se admiten archivos PDF y DOCX.",
                        },
                    )
                except Exception:
                    pass

            return {
                "ok": False,
                "detail": "Formato no permitido. Solo se admiten archivos PDF y DOCX."
            }

        if is_pdf:
            # Si el PDF tiene muchas páginas, no se rechaza: se procesa hasta el límite blando.
            pdf_pages = self._count_pdf_pages_bytes(data)
            text = self._extract_text_pdf_bytes(data, max_pages=MAX_PDF_PAGES)

        elif is_docx:
            text = self._extract_text_docx_bytes(data)

        # Solo se rechaza si no hay texto útil real; si hay demasiado, se recorta.
        text = self._clean_text(text)

        if not text or len(text) < MIN_TEXT_CHARS:
            if self.espacio_archivos_repo:
                try:
                    await self.espacio_archivos_repo.actualizar(
                        access_token,
                        file_id,
                        {
                            "estado": "failed",
                            "error_detail": (
                                "No se detectó suficiente texto utilizable en el archivo. "
                                "Sube un PDF o DOCX con contenido textual real, no escaneado ni basado principalmente en imágenes."
                            )[:1000],
                        },
                    )
                except Exception:
                    pass

            return {
                "ok": False,
                "detail": (
                    "No se detectó suficiente texto utilizable en el archivo. "
                    "Sube un PDF o DOCX con contenido textual real, no escaneado ni basado principalmente en imágenes."
                ),
            }

        if len(text) > MAX_TEXT_CHARS:
            print(
                f"[INGESTA] texto recortado: original={len(text)} "
                f"limite={MAX_TEXT_CHARS} archivo={filename}"
            )
            text = text[:MAX_TEXT_CHARS]
        
        
        t0 = time.time()
        print(f"[INGESTA] archivo={filename} ctype={content_type} bytes={len(data)} text_len={len(text)}")

        chunks = self._chunk_text(text, max_chars=900, overlap=120)
        print(f"[INGESTA] chunks={len(chunks)} chunking_s={time.time()-t0:.2f}")


        if not chunks:
            if self.espacio_archivos_repo:
                try:
                    await self.espacio_archivos_repo.actualizar(
                        access_token,
                        file_id,
                        {
                            "estado": "failed",
                            "error_detail": "No se generaron chunks",
                        },
                    )
                except Exception:
                    pass

            return {"ok": False, "detail": "No se generaron chunks"}

        # Si genera demasiados chunks, no se rechaza: se procesa solo el máximo permitido.
        if len(chunks) > MAX_CHUNKS:
            original_chunks = len(chunks)
            chunks = chunks[:MAX_CHUNKS]
            print(
                f"[INGESTA] chunks recortados: original={original_chunks} "
                f"limite={MAX_CHUNKS} archivo={filename}"
            )
        
        
        # Se preparan identificadores y parámetros de procesamiento por lotes para hacer
        # la inserción de embeddings y etiquetas de forma más eficiente.
        inserted = 0
        t_loop = time.time()
        print(f"[INGESTA] start_insert chunks={len(chunks)}")

        # Tamaño de los lotes (batch): cantidad de registros que se acumulan antes
        # de insertarlos juntos en la base de datos para reducir número de requests.
        # (ajustar si es necesario según rendimiento y límites de tu infraestructura)
        BATCH_EMB = 50
        BATCH_RED2 = 200  # red2 payloads son livianos, puede ser mayor


        # A partir de aquí se recorre cada chunk del archivo y se va construyendo toda
        # la información necesaria para embeddings, RED1 y RED2 antes de hacer flush por lotes.
        emb_payloads: List[Dict[str, Any]] = []
        red2_payloads: List[Dict[str, Any]] = []

        # cache por chunk para no recalcular
        red1_cache: Dict[int, Dict[str, Any]] = {}

        total = len(chunks)

        for idx, chunk in enumerate(chunks):
            if idx % 25 == 0:
                print(f"[INGESTA] idx={idx}/{total} inserted={inserted} elapsed_s={time.time()-t_loop:.1f}")

            # 1) Para cada chunk se genera su embedding y se acumula el payload con sus metadatos
            # para insertarlo más adelante en un solo lote.
            vec = self.embeddings_model.embed(chunk)
            emb_payloads.append(
                {
                    "docente_id": docente_id,
                    "espacio_id": espacio_id,
                    "tipo_fuente": "archivo",
                    "fuente_id": file_id,
                    "texto": chunk,
                    "embedding": vec,
                    "metadatos": {
                        "ambito": "espacio",
                        "espacio_id": espacio_id,
                        "archivo": filename,
                        "fuente_id": file_id,
                        "file_hash": file_hash,
                        "chunk_index": idx,
                        "total_chunks": total,
                        "content_type": content_type,
                        "espacio_nombre": espacio.get("nombre"),
                    },
                }
            )

            # 2) Luego se intenta clasificar el chunk con RED1 y guardar sus resultados. Además,
            # se conservan las salidas necesarias para enriquecer posteriormente a RED2.
            red1_areas_top = None
            red1_dims_probs = None
            try:
                red1_pack = await self.red1.clasificar_y_guardar(
                    access_token,
                    docente_id=docente_id,
                    espacio_id=espacio_id,
                    conversacion_espacio_id=None,
                    mensaje_espacio_id=None,
                    tipo_fuente="archivo",
                    fuente_id=file_id,
                    chunk_index=idx,
                    texto=chunk,
                )
                if isinstance(red1_pack, dict):
                    red1_areas_top = red1_pack.get("areas_top")
                    red1_dims_probs = red1_pack.get("dims_probs")
            except Exception:
                pass

            red1_cache[idx] = {
                "areas_top": red1_areas_top,
                "dims_probs": red1_dims_probs,
            }

            # 3) Cuando el lote alcanza el tamaño definido o cuando ya se llegó al último chunk,
            # se inserta el batch completo de embeddings en la base de datos.
            is_last = (idx == total - 1)
            should_flush_emb = (len(emb_payloads) >= BATCH_EMB) or is_last

            if should_flush_emb:
                t_ins = time.time()

                # La inserción por lotes reduce la cantidad de requests y mejora bastante el
                # rendimiento frente a insertar cada chunk de manera individual. 1 request para N filas
                inserted_rows = await self.embeddings_repo.insert_many_embeddings(access_token, emb_payloads)
                inserted += len(inserted_rows)


                # Después del insert se reconstruye la relación entre cada chunk y su embedding
                # persistido para que RED2 pueda guardar sus resultados apuntando al id correcto.
                id_by_chunk: Dict[int, str] = {}
                for r in inserted_rows:
                    meta = r.get("metadatos") or {}
                    ci = meta.get("chunk_index")
                    if isinstance(ci, int) and r.get("id"):
                        id_by_chunk[ci] = r["id"]

                print(f"[INGESTA] embeddings_batch n={len(inserted_rows)} batch_s={time.time()-t_ins:.2f}")

                # Si RED2 está disponible, se generan sus predicciones tomando como base el texto
                # del chunk, las señales de RED1 y el id real del embedding recién insertado.
                if self.red2_model and self.red2_repo:
                    for r in inserted_rows:
                        meta = r.get("metadatos") or {}
                        ci = meta.get("chunk_index")
                        if not isinstance(ci, int):
                            continue

                        chunk_text = r.get("texto") or ""
                        if not chunk_text:
                            continue

                        red1_areas_top = (red1_cache.get(ci) or {}).get("areas_top")
                        red1_dims_probs = (red1_cache.get(ci) or {}).get("dims_probs")

                        try:
                            out2 = self.red2_model.predict(
                                text=chunk_text,
                                areas_top=red1_areas_top,
                                dims_probs=red1_dims_probs,
                                tipo_fuente="archivo",
                                top_k=5,
                                query_for_post=None,
                            )

                            red2_payloads.append(
                                {
                                    "docente_id": docente_id,
                                    "espacio_id": espacio_id,
                                    "embedding_texto_id": r["id"], 
                                    "fuente_id": file_id,
                                    "tipo_fuente": "archivo",
                                    "red1_areas_top": red1_areas_top,
                                    "red1_dims_probs": red1_dims_probs,
                                    "red2_top": out2.top,
                                    "red2_probs": out2.probs,
                                    "metadatos": {
                                        "archivo": filename,
                                        "chunk_index": ci,
                                        "total_chunks": total,
                                    },
                                }
                            )
                        except Exception:
                            pass

                    # Las etiquetas de RED2 también se insertan en lote para no degradar el tiempo
                    # total de ingesta cuando el archivo genera muchos chunks.
                    if len(red2_payloads) >= BATCH_RED2 or is_last:
                        t_r2 = time.time()
                        try:
                            await self.red2_repo.insert_chunk_labels_many(access_token, red2_payloads)
                        except Exception:
                            pass
                        print(f"[INGESTA] red2_batch n={len(red2_payloads)} batch_s={time.time()-t_r2:.2f}")
                        red2_payloads.clear()

                # limpiar batch embeddings
                emb_payloads.clear()

 
        # Al finalizar la ingesta se registra el evento en RED3 y se intenta actualizar
        # el perfil del docente con esta nueva actividad, sin romper el flujo si falla.
        if hasattr(self, "red3_service") and self.red3_service:
            try:
                await self.red3_service.record_event_best_effort(
                    access_token,
                    docente_id=docente_id,
                    event_type="espacio_file_upload",
                    meta={
                        "kind": "espacio_archivo",
                        "espacio_id": espacio_id,
                        "fuente_id": file_id,
                        "archivo": filename,
                        "chunks": inserted,
                        "file_hash": file_hash,
                        "content_type": content_type,
                    },
                )

                await self.red3_service.update_profile_best_effort(
                    access_token,
                    docente_id=docente_id,
                    window_days=30,
                )
            except Exception:
                pass

        print(f"[INGESTA] done inserted={inserted} total_s={time.time()-t0:.1f}")

        # Si el archivo terminó de procesarse correctamente, se actualiza su registro maestro.
        if self.espacio_archivos_repo:
            try:
                await self.espacio_archivos_repo.actualizar(
                    access_token,
                    file_id,
                    {
                        "estado": "processed",
                        "total_chunks": inserted,
                        "error_detail": None,
                    },
                )
            except Exception:
                pass

        return {
            "ok": True,
            "espacio_id": espacio_id,
            "docente_id": docente_id,
            "archivo": filename,
            "file_id": file_id,
            "chunks_insertados": inserted,
        }


    # Limpia caracteres problemáticos y normaliza espacios y saltos de línea para
    # dejar el texto en mejores condiciones antes de fragmentarlo.
    def _clean_text(self, t: str) -> str:
        t = t.replace("\x00", " ")
        t = re.sub(r"[ \t]+", " ", t)
        t = re.sub(r"\n{3,}", "\n\n", t)
        return t.strip()

    # Divide el texto en chunks relativamente cortos con solapamiento controlado
    # para conservar algo de contexto entre fragmentos consecutivos.
    def _chunk_text(self, t: str, max_chars: int = 900, overlap: int = 120) -> List[str]:
        paras = [p.strip() for p in t.split("\n") if p.strip()]
        chunks: List[str] = []
        buf = ""

        # Esta función auxiliar cierra el buffer actual y lo agrega a la lista final
        # de chunks cuando ya no entra más contenido.
        def flush():
            nonlocal buf
            if buf.strip():
                chunks.append(buf.strip())
            buf = ""

        for p in paras:
            if len(buf) + len(p) + 1 <= max_chars:
                buf = (buf + "\n" + p).strip() if buf else p
            else:
                flush()
                if len(p) <= max_chars:
                    buf = p
                else:
                    start = 0
                    while start < len(p):
                        end = min(start + max_chars, len(p))
                        chunks.append(p[start:end].strip())

                        if end >= len(p):  # Evita loop infinito al llegar al final
                            break

                        start = max(0, end - overlap)
        flush()

        # Aquí se vuelve a inyectar una pequeña cola del chunk anterior en el siguiente
        # para no perder continuidad semántica entre fragmentos consecutivos.
        if overlap > 0 and len(chunks) > 1:
            out = []
            for i, c in enumerate(chunks):
                if i == 0:
                    out.append(c)
                else:
                    prev = chunks[i - 1]
                    tail = prev[-overlap:] if len(prev) > overlap else prev
                    out.append((tail + "\n" + c).strip())
            chunks = out

        return [c for c in chunks if len(c) >= 30]


    # Cuenta páginas de un PDF para frenar documentos demasiado extensos antes de procesarlos.
    def _count_pdf_pages_bytes(self, data: bytes) -> int:
        try:
            from pypdf import PdfReader
            import io
            reader = PdfReader(io.BytesIO(data))
            return len(reader.pages)
        except Exception:
            return 0

    # Extrae texto plano desde un PDF cargado en memoria usando pypdf.
    # Si el documento tiene muchas páginas, solo procesa hasta el límite configurado.
    def _extract_text_pdf_bytes(self, data: bytes, max_pages: int | None = None) -> str:
        try:
            from pypdf import PdfReader
            import io
            reader = PdfReader(io.BytesIO(data))
            parts = []

            pages = reader.pages
            if isinstance(max_pages, int) and max_pages > 0:
                pages = reader.pages[:max_pages]

            for page in pages:
                parts.append(page.extract_text() or "")

            return "\n".join(parts)
        except Exception:
            return ""
        

    # Extrae texto desde un archivo DOCX cargado en memoria, leyendo tanto párrafos
    # como tablas para recuperar la mayor cantidad posible de contenido útil.
    def _extract_text_docx_bytes(self, data: bytes) -> str:
        try:
            import io
            from docx import Document  # python-docx

            doc = Document(io.BytesIO(data))
            parts: list[str] = []

            # Párrafos
            for p in doc.paragraphs:
                if p.text and p.text.strip():
                    parts.append(p.text.strip())

            # Tablas (si tienen contenido)
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))

            return "\n".join(parts)
        except Exception:
            return ""