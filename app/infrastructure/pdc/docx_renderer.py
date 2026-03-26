# app/infrastructure/pdc/docx_renderer.py

# Módulo de infraestructura encargado de construir el documento final del PDC
# en formato DOCX. Aquí se definen helpers de estilo y la función principal
# que toma el payload original y el contenido generado para renderizar el
# documento Word con la estructura formal requerida por el sistema.

from io import BytesIO
from typing import Any, Dict, List, Union

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENT


# Helpers de estilo
# Configura los márgenes generales del documento para mantener un formato
# uniforme y más cercano al estilo final esperado del PDC.
def _set_page(doc: Document):
    sec = doc.sections[0]

    # ORIENTACIÓN HORIZONTAL
    sec.orientation = WD_ORIENT.LANDSCAPE

    # Intercambiar ancho y alto (OBLIGATORIO en python-docx)
    sec.page_width, sec.page_height = sec.page_height, sec.page_width

    # Márgenes
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin = Inches(0.6)
    sec.right_margin = Inches(0.6)

# Helper de formato básico para insertar texto dentro de un párrafo con una
# tipografía y tamaño consistentes en todo el documento.
def _set_run(p, text, bold=False, size=11):
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Calibri"
    return r

# Aplica color de fondo a una celda de tabla usando manipulación XML, útil
# para resaltar encabezados o secciones importantes.
#def _shade_cell(cell, fill_hex: str):
 #   tcPr = cell._tc.get_or_add_tcPr()
 #   shd = OxmlElement("w:shd")
 #   shd.set(qn("w:val"), "clear")
 #   shd.set(qn("w:color"), "auto")
 #   shd.set(qn("w:fill"), fill_hex)
 #   tcPr.append(shd)

# Aplica bordes visibles a una celda individual para mantener el estilo visual
# de tablas y cajas dentro del documento.
def _set_cell_borders(cell, color="000000", size="8"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        element = tcBorders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tcBorders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)

# Aplica bordes a toda una tabla, incluyendo bordes internos, para que su
# estructura quede claramente delimitada en el documento final.
def _set_table_borders(table, color="000000", size="8"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tblBorders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tblBorders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)

# Limpia el contenido inicial de una celda y devuelve su primer párrafo para
# poder reutilizarla con formato controlado.
def _clear_cell(cell):
    cell.text = ""
    return cell.paragraphs[0]

# Normaliza distintos tipos de entrada a una lista de strings para facilitar
# el render uniforme de listas, textos simples o valores sueltos.
def _as_list(v: Any) -> List[str]:
    if v is None:
        return []
    if isinstance(v, list):
        return [str(x).strip() for x in v if str(x).strip()]
    if isinstance(v, str):
        s = v.strip()
        if not s:
            return []
        if "\n" in s:
            return [x.strip("•- \t") for x in s.split("\n") if x.strip()]
        return [s]
    return [str(v).strip()]

def _text_or_join(value: Union[str, List[str], None]) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        return "\n".join(str(x).strip() for x in value if str(x).strip())
    return str(value).strip()


def _merge_row(table, row_idx: int, start_col: int, end_col: int):
    cell = table.cell(row_idx, start_col)
    for col in range(start_col + 1, end_col + 1):
        cell = cell.merge(table.cell(row_idx, col))
    return cell


def _write_cell(cell, text: str, bold: bool = False, size: int = 10, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = _clear_cell(cell)
    p.alignment = align
    _set_run(p, text or "", bold=bold, size=size)
    return p


def _write_multiline_cell(cell, lines: Union[str, List[str]], bold_first: bool = False, size: int = 10):
    p = _clear_cell(cell)

    if isinstance(lines, list):
        clean = [str(x).strip() for x in lines if str(x).strip()]
    else:
        raw = str(lines or "").strip()
        clean = [ln.strip() for ln in raw.split("\n") if ln.strip()]

    if not clean:
        _set_run(p, "", size=size)
        return

    for i, ln in enumerate(clean):
        px = p if i == 0 else cell.add_paragraph()
        _set_run(px, ln, bold=(bold_first and i == 0), size=size)
        px.paragraph_format.space_after = Pt(1)


def _criterio_to_text(v: Any) -> str:
    if isinstance(v, list):
        return " ".join(str(x).strip() for x in v if str(x).strip())
    return str(v or "").strip()


def _to_int(v: Any, default: int = 0) -> int:
    try:
        return int(str(v).strip())
    except Exception:
        return default


def _periodos_label_from_ident(ident: Dict[str, Any]) -> str:
    semanas = _to_int(ident.get("semanas"), 0)
    periodos_por_semana = _to_int(ident.get("periodos_por_semana"), 0)
    total = semanas * periodos_por_semana if semanas > 0 and periodos_por_semana > 0 else 0
    return f"{total} periodos" if total > 0 else str(ident.get("tiempo", "") or "").strip()


def _split_week_contents(contenidos: List[str]) -> List[str]:
    out: List[str] = []
    for item in contenidos:
        s = str(item or "").strip()
        if not s:
            continue
        out.append(s)
    return out


def _prefix_lines(title: str, value: Union[str, List[str], None]) -> List[str]:
    lines = _as_list(value)
    if not lines:
        return [f"{title}:"]
    return [f"{title}:"] + [f"• {x}" for x in lines]



# Función principal que arma el documento DOCX del PDC. Toma los datos de
# entrada del formulario (`payload`) y el contenido generado por el sistema
# (`generado`), organiza las secciones del documento y devuelve el archivo
# final en memoria listo para descarga o almacenamiento.
def render_pdc_docx(payload: dict, generado: dict) -> BytesIO:
    # Inicializa el documento Word y aplica la configuración base de página.
    doc = Document()
    _set_page(doc)

    ident = payload.get("identificacion", {}) or {}
    contexto = payload.get("contexto", {}) or {}
    variables = payload.get("variables", {}) or {}
    contenidos = variables.get("contenidos", []) or []

    use_psp = bool(contexto.get("use_psp", True))

    unidad_educativa = ident.get("unidad_educativa", "") or ""
    area = ident.get("area", "") or ""
    nivel = ident.get("nivel", "") or ""
    anio = ident.get("anio_escolaridad", "") or ""
    trimestre = ident.get("trimestre", "") or ""
    docente = ident.get("docente", "") or ""

    semanas = _to_int(ident.get("semanas"), 0)
    periodos_por_semana = _to_int(ident.get("periodos_por_semana"), 0)
    duracion_periodo = str(ident.get("duracion_periodo", "") or "").strip()
    tiempo = _periodos_label_from_ident(ident)

    # En la plantilla final se mostrará objetivo de aprendizaje.
    # Con PSP usamos el objetivo generado porque articula PAT + PSP.
    # Sin PSP usamos el input directo del docente.
    objetivo_aprendizaje = (
        str(generado.get("objetivo_holistico", "") or "").strip()
        if use_psp
        else str(contexto.get("objetivo_aprendizaje", "") or "").strip()
    )

    practica = _text_or_join(generado.get("practica"))
    teoria = _text_or_join(generado.get("teoria"))
    valoracion = _text_or_join(generado.get("valoracion"))
    produccion = _text_or_join(generado.get("produccion"))
    recursos = _text_or_join(generado.get("recursos"))
    producto = _text_or_join(generado.get("producto"))

    criterios = generado.get("criterios", {}) or {}
    ser = _criterio_to_text(criterios.get("SER"))
    saber = _criterio_to_text(criterios.get("SABER"))
    hacer = _criterio_to_text(criterios.get("HACER"))
    decidir = _criterio_to_text(criterios.get("DECIDIR"))

    # Encabezado principal
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("PLAN DE DESARROLLO CURRICULAR")
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Calibri"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("SECUNDARIA COMUNITARIA PRODUCTIVA")
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.name = "Calibri"

    doc.add_paragraph("")

    # DATOS REFERENCIALES
    distrito_educativo = str(ident.get("distrito_educativo", "") or "").strip()
    fecha_inicio = str(ident.get("fecha_inicio", "") or "").strip()
    fecha_fin = str(ident.get("fecha_fin", "") or "").strip()

    t1 = doc.add_table(rows=8, cols=4)
    t1.autofit = True
    _set_table_borders(t1)

    # Encabezado
    c = _merge_row(t1, 0, 0, 3)
    _write_cell(c, "DATOS REFERENCIALES", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Fila 1
    _write_cell(t1.cell(1, 0), "Distrito educativo", bold=True)
    _write_cell(t1.cell(1, 1), distrito_educativo)
    _write_cell(t1.cell(1, 2), "Unidad Educativa", bold=True)
    _write_cell(t1.cell(1, 3), unidad_educativa)

    # Fila 2
    _write_cell(t1.cell(2, 0), "Nivel", bold=True)
    _write_cell(t1.cell(2, 1), nivel)
    _write_cell(t1.cell(2, 2), "Año de escolaridad/\nParalelo", bold=True)
    _write_cell(t1.cell(2, 3), anio)

    # Fila 3
    _write_cell(t1.cell(3, 0), "Maestra/o", bold=True)
    c = _merge_row(t1, 3, 1, 3)
    _write_cell(c, docente)

    # Fila 4
    _write_cell(t1.cell(4, 0), "Área", bold=True)
    c = _merge_row(t1, 4, 1, 3)
    _write_cell(c, area)

    # Fila 5
    _write_cell(t1.cell(5, 0), "Trimestre", bold=True)
    c = _merge_row(t1, 5, 1, 3)
    _write_cell(c, trimestre)

    # Fila 6
    _write_cell(t1.cell(6, 0), "Fecha", bold=True)
    c = _merge_row(t1, 6, 1, 3)
    fecha_text = ""
    if fecha_inicio and fecha_fin:
        fecha_text = f"Del: {fecha_inicio}        Al: {fecha_fin}"
    elif fecha_inicio:
        fecha_text = f"Del: {fecha_inicio}"
    elif fecha_fin:
        fecha_text = f"Al: {fecha_fin}"
    _write_cell(c, fecha_text)

    # Fila 7
    _write_cell(t1.cell(7, 0), "Tiempo", bold=True)
    c = _merge_row(t1, 7, 1, 3)
    _write_cell(c, tiempo)

    doc.add_paragraph("")

    # =========================================================
    # TABLA PRINCIPAL (FORMATO FISCAL REAL)
    # =========================================================
    t2 = doc.add_table(rows=2, cols=6)
    t2.style = "Table Grid"

    headers = [
        "OBJETIVO DE APRENDIZAJE",
        f"CONTENIDOS ({semanas} semanas)" if semanas else "CONTENIDOS",
        "MOMENTOS DEL PROCESO FORMATIVO",
        "RECURSOS",
        "PERÍODOS",
        "CRITERIOS DE EVALUACIÓN (SER – SABER – HACER – DECIDIR)",
    ]

    # Header row
    for i, h in enumerate(headers):
        _write_cell(t2.cell(0, i), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Contenido
    t2.cell(1, 0).text = objetivo_aprendizaje

    # CONTENIDOS
    _write_multiline_cell(
        t2.cell(1, 1),
        _split_week_contents(contenidos),
        size=10,
    )

    # MOMENTOS
    momentos = (
        _prefix_lines("Práctica", generado.get("practica"))
        + _prefix_lines("Teoría", generado.get("teoria"))
        + _prefix_lines("Valoración", generado.get("valoracion"))
        + _prefix_lines("Producción", generado.get("produccion"))
    )

    _write_multiline_cell(
        t2.cell(1, 2),
        momentos,
        size=10,
    )

    # RECURSOS
    _write_multiline_cell(
        t2.cell(1, 3),
        _as_list(generado.get("recursos")),
        size=10,
    )

    # PERIODOS
    periodos_text = tiempo
    if semanas > 0 and periodos_por_semana > 0 and duracion_periodo:
        periodos_text = f"{tiempo} ({periodos_por_semana} por semana, {duracion_periodo} min c/u)"

    periodos_lines = []

    if semanas > 0 and periodos_por_semana > 0:
        for i in range(1, semanas + 1):
            periodos_lines.append(f"Semana {i}: {periodos_por_semana} periodos")

        if duracion_periodo:
            periodos_lines.append(f"Duración: {duracion_periodo} min por periodo")
    else:
        periodos_lines.append(periodos_text)

    _write_multiline_cell(
        t2.cell(1, 4),
        periodos_lines,
        size=10,
    )

    # CRITERIOS
    criterios_lines = [
        f"SER: {_criterio_to_text(criterios.get('SER'))}",
        f"SABER: {_criterio_to_text(criterios.get('SABER'))}",
        f"HACER: {_criterio_to_text(criterios.get('HACER'))}",
        f"DECIDIR: {_criterio_to_text(criterios.get('DECIDIR'))}",
    ]

    _write_multiline_cell(
        t2.cell(1, 5),
        criterios_lines,
        size=10,
    )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf