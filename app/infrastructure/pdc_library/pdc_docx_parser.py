# app/infrastructure/pdc_library/pdc_docx_parser.py
from __future__ import annotations

from io import BytesIO
from typing import Any, Dict, List, Tuple
import re

from docx import Document


# =========================
# Normalización
# =========================
def _norm(s: str) -> str:
    s = (s or "").replace("\u00a0", " ")
    # NO destruye símbolos (✓, etc). Solo colapsa espacios.
    return re.sub(r"\s+", " ", s).strip()


_BULLET_RE = re.compile(r"^\s*(?:✓|•|\-|\*|\d+[\.\)]|[a-zA-Z][\.\)])\s+")


def _split_to_list(text: str) -> List[str]:
    """
    Convierte texto a lista sin inventar splits raros:
    - Si hay viñetas/numeración -> lista por líneas
    - Si no -> un solo párrafo
    """
    t = (text or "").strip()
    if not t:
        return []

    lines = [ln.strip() for ln in t.splitlines() if ln.strip()]
    if len(lines) <= 1:
        return [_norm(t)]

    bullet_like = sum(1 for ln in lines if _BULLET_RE.match(ln))
    if bullet_like >= max(1, len(lines) // 4):
        out: List[str] = []
        for ln in lines:
            ln2 = _BULLET_RE.sub("", ln).strip()
            if ln2:
                out.append(_norm(ln2))
        return out if out else [_norm(t)]

    return [_norm(t)]


# =========================
# Extracción docx: párrafos + celdas
# =========================
def _iter_doc_blocks(doc: Document) -> List[Tuple[str, str]]:
    """
    Devuelve lista de bloques:
      ('p', 'texto') para paragraphs
      ('cell', 'texto\\ntexto') para celdas de tablas
    """
    out: List[Tuple[str, str]] = []

    for p in doc.paragraphs:
        tx = _norm(p.text)
        if tx:
            out.append(("p", tx))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_texts = []
                for p in cell.paragraphs:
                    tx = _norm(p.text)
                    if tx:
                        cell_texts.append(tx)
                if cell_texts:
                    out.append(("cell", "\n".join(cell_texts)))

    return out


# =========================
# Parser ministerial robusto
# =========================
# OJO: aquí NO hacemos match de "teoría" suelta, sino de marcadores "TEORÍA ✓"
_SEC_RE = re.compile(
    r"(?i)\b(PRÁCTICA|PRACTICA|TEORÍA|TEORIA|VALORACIÓN|VALORACION|PRODUCCIÓN|PRODUCCION)\b\s*(?:✓|:|-)\s*"
)

_DIM_RE = re.compile(r"(?i)\b(SER|SABER|HACER|DECIDIR)\b\s*(?:✓|:|-)\s*")


def _split_orientaciones(text: str) -> Dict[str, str]:
    """
    Corta una celda tipo:
      "PRÁCTICA ✓ ... TEORÍA ✓ ... VALORACIÓN ✓ ... PRODUCCIÓN ✓ ..."
    usando marcadores fuertes (palabra + ✓/:/-).
    """
    txt = text or ""
    ms = list(_SEC_RE.finditer(txt))
    if not ms:
        return {}

    parts: Dict[str, str] = {}
    for i, m in enumerate(ms):
        key = m.group(1).upper()
        start = m.end()
        end = ms[i + 1].start() if i + 1 < len(ms) else len(txt)
        seg = txt[start:end].strip()
        if seg:
            parts[key] = seg
    return parts


def _split_criterios(text: str) -> Dict[str, str]:
    """
    Corta texto tipo:
      "SER ✓ ... SABER ✓ ... HACER ✓ ... DECIDIR ✓ ..."
    """
    out = {"SER": "", "SABER": "", "HACER": "", "DECIDIR": ""}
    t = text or ""
    ms = list(_DIM_RE.finditer(t))
    if not ms:
        return out

    for i, m in enumerate(ms):
        dim = m.group(1).upper()
        start = m.end()
        end = ms[i + 1].start() if i + 1 < len(ms) else len(t)
        seg = t[start:end].strip()
        seg = re.sub(r"^\s*✓\s*", "", seg).strip()
        out[dim] = _norm(seg)
    return out


def parse_pdc_docx_to_red3_bloques(docx_bytes: bytes) -> Dict[str, Any]:
    empty = {
        "document": {
            "text_preview": "",
            "text_length": 0,
        },
        "blocks": {
            "main_pedagogical_block": "",
            "evaluation_product_block": "",
            "resources_time_block": "",
        },
        "analysis": {
            "has_product_keyword": False,
            "has_resources_keyword": False,
            "has_eval_keyword": False,
            "has_week_structure": False,
        },
    }

    try:
        doc = Document(BytesIO(docx_bytes))
    except Exception:
        return empty

    blocks = _iter_doc_blocks(doc)

    all_texts: List[str] = []
    for _, tx in blocks:
        t = _norm(tx)
        if t:
            all_texts.append(t)

    def is_referential(p: str) -> bool:
        pl = p.lower()
        return any(k in pl for k in [
            "plan de desarrollo curricular",
            "datos referenciales",
            "distrito educativo",
            "unidad educativa",
            "nivel",
            "año de escolaridad",
            "anio de escolaridad",
            "paralelo",
            "docente",
            "maestra",
            "maestro",
            "trimestre",
            "fecha",
        ])

    filtered_texts = [p for p in all_texts if not is_referential(p)]

    text_full = " ".join(filtered_texts).strip()
    text_preview = text_full[:500]

    main_block_parts: List[str] = []
    eval_block_parts: List[str] = []
    resources_block_parts: List[str] = []

    for p in filtered_texts:
        pl = p.lower()

        if any(k in pl for k in ["criterio", "evaluación", "evaluacion", "ser:", "saber:", "hacer:", "decidir:", "producto", "productos:"]):
            eval_block_parts.append(p)
        elif any(k in pl for k in ["recurso", "periodo", "período", "semana", "duración", "duracion", "tiempo"]):
            resources_block_parts.append(p)
        else:
            main_block_parts.append(p)

    main_block = " ".join(main_block_parts).strip()
    eval_block = " ".join(eval_block_parts).strip()
    resources_block = " ".join(resources_block_parts).strip()

    has_week_structure = any("semana" in p.lower() for p in filtered_texts)
    has_product_keyword = any("producto" in p.lower() for p in filtered_texts)
    has_resources_keyword = any("recurso" in p.lower() for p in filtered_texts)
    has_eval_keyword = any(k in text_full.lower() for k in ["criterio", "evaluación", "evaluacion", "ser:", "saber:", "hacer:", "decidir:"])

    return {
        "document": {
            "text_preview": text_preview,
            "text_length": len(text_full),
        },
        "blocks": {
            "main_pedagogical_block": main_block,
            "evaluation_product_block": eval_block,
            "resources_time_block": resources_block,
        },
        "analysis": {
            "has_product_keyword": has_product_keyword,
            "has_resources_keyword": has_resources_keyword,
            "has_eval_keyword": has_eval_keyword,
            "has_week_structure": has_week_structure,
        },
    }